from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

import boto3
import fitz
from botocore.config import Config
from docx import Document
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel, Field


app = FastAPI(title="YuDu_Bidkit Document Worker")


class StorageConfig(BaseModel):
    endpoint: str
    bucket: str
    region: str = "auto"
    access_key_id: str = Field(alias="accessKeyId")
    secret_access_key: str = Field(alias="secretAccessKey")
    force_path_style: bool = Field(default=True, alias="forcePathStyle")
    use_ssl: bool = Field(default=False, alias="useSSL")


class ParseRequest(BaseModel):
    storage_key: str
    parser: str = "auto"
    file_name: str | None = None
    mime_type: str | None = None
    storage: StorageConfig


def s3_client(config: StorageConfig):
    return boto3.client(
        "s3",
        endpoint_url=config.endpoint.rstrip("/"),
        region_name=config.region,
        aws_access_key_id=config.access_key_id,
        aws_secret_access_key=config.secret_access_key,
        use_ssl=config.use_ssl,
        config=Config(s3={"addressing_style": "path" if config.force_path_style else "virtual"}),
    )


def extension_for(request: ParseRequest) -> str:
    name = request.file_name or request.storage_key
    return Path(name).suffix.lower()


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            level = "".join(character for character in style_name if character.isdigit()) or "2"
            chunks.append(f"{'#' * min(max(int(level), 1), 6)} {text}")
        else:
            chunks.append(text)

    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = normalized[0]
        chunks.append("| " + " | ".join(header) + " |")
        chunks.append("| " + " | ".join(["---"] * width) + " |")
        for row in normalized[1:]:
            chunks.append("| " + " | ".join(row) + " |")

    return "\n\n".join(chunks).strip()


def parse_pdf(data: bytes) -> str:
    document = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for index, page in enumerate(document, start=1):
        text = page.get_text("text").strip()
        if text:
            pages.append(f"## 第 {index} 页\n\n{text}")
    return "\n\n".join(pages).strip()


def parse_to_markdown(request: ParseRequest, data: bytes) -> tuple[str, str]:
    extension = extension_for(request)
    parser = request.parser

    if parser == "auto":
        if extension == ".docx":
            parser = "docx"
        elif extension == ".pdf":
            parser = "pdf"
        elif extension in {".md", ".markdown", ".txt"}:
            parser = "text"
        else:
            parser = "text"

    if parser == "docx":
        return parse_docx(data), parser
    if parser == "pdf":
        return parse_pdf(data), parser
    if parser in {"text", "md", "markdown"}:
        return decode_text(data), "text"

    raise HTTPException(status_code=400, detail=f"暂不支持解析器：{request.parser}")


def markdown_key_for(storage_key: str) -> str:
    path = Path(storage_key)
    stem = path.name.rsplit(".", 1)[0] if "." in path.name else path.name
    parent = str(path.parent).replace("\\", "/")
    if parent == ".":
        parent = "uploads"
    return f"{parent.replace('/documents/', '/markdown/')}/{stem}.md"


def markdown_stats(markdown: str) -> dict[str, Any]:
    lines = markdown.splitlines()
    headings = [line.lstrip("#").strip() for line in lines if line.startswith("#")]
    return {
        "char_count": len(markdown),
        "line_count": len(lines),
        "heading_count": len(headings),
        "headings": headings[:20],
    }


@app.get("/health")
def health():
    return {"status": "ok", "service": "document-worker"}


@app.post("/parse")
def parse_document(request: ParseRequest):
    client = s3_client(request.storage)

    try:
        obj = client.get_object(Bucket=request.storage.bucket, Key=request.storage_key)
        raw = obj["Body"].read()
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=502, detail=f"读取对象存储失败：{exc}") from exc

    markdown, parser = parse_to_markdown(request, raw)
    markdown_key = markdown_key_for(request.storage_key)

    try:
        client.put_object(
            Bucket=request.storage.bucket,
            Key=markdown_key,
            Body=markdown.encode("utf-8"),
            ContentType="text/markdown; charset=utf-8",
        )
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=502, detail=f"写入 Markdown 失败：{exc}") from exc

    return {
        "status": "parsed",
        "storage_key": request.storage_key,
        "markdown_key": markdown_key,
        "parser": parser,
        "markdown": markdown[:8000],
        "stats": markdown_stats(markdown),
        "message": "文档解析完成，Markdown 中间态已写入对象存储。",
    }
